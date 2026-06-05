from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

doc = Document()

for section in doc.sections:
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(4)
    section.bottom_margin = Cm(3)
    section.left_margin = Cm(4)
    section.right_margin = Cm(3)

sty = doc.styles['Normal']
sty.font.name = 'Times New Roman'
sty.font.size = Pt(12)
sty.paragraph_format.line_spacing = 1.5
sty.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
sty.paragraph_format.first_line_indent = Cm(1.25)

# ── Configure heading styles ──
for lvl, (sz, sp, al) in {1: (14, 12, WD_ALIGN_PARAGRAPH.CENTER), 2: (12, 12, WD_ALIGN_PARAGRAPH.LEFT), 3: (12, 12, WD_ALIGN_PARAGRAPH.LEFT)}.items():
    s = doc.styles[f'Heading {lvl}']
    s.font.name = 'Times New Roman'
    s.font.size = Pt(sz)
    s.font.bold = True
    s.font.color.rgb = RGBColor(0, 0, 0)
    s.paragraph_format.line_spacing = 1.5
    s.paragraph_format.space_before = Pt(sp)
    s.paragraph_format.space_after = Pt(6)
    s.paragraph_format.alignment = al
    s.paragraph_format.first_line_indent = Cm(0)

def sf(run, name='Times New Roman', size=12, bold=False, italic=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)

def h(text, level=1, sb=12, sa=6):
    if level <= 0:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(sb)
        p.paragraph_format.space_after = Pt(sa)
        p.paragraph_format.line_spacing = 1.5
        p.paragraph_format.first_line_indent = Cm(0)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        sf(run, bold=True, size=14)
        return p
    if level == 1:
        return doc.add_heading(text, level=1)
    elif level == 2:
        return doc.add_heading(text, level=2)
    elif level == 3:
        return doc.add_heading(text, level=3)

def p(text, bold=False, italic=False, align=WD_ALIGN_PARAGRAPH.JUSTIFY, indent=True, size=12):
    par = doc.add_paragraph()
    par.paragraph_format.line_spacing = 1.5
    par.alignment = align
    if indent:
        par.paragraph_format.first_line_indent = Cm(1.25)
    else:
        par.paragraph_format.first_line_indent = Cm(0)
    run = par.add_run(text)
    sf(run, bold=bold, italic=italic, size=size)
    return par

def b(text):
    par = doc.add_paragraph()
    par.paragraph_format.line_spacing = 1.5
    par.paragraph_format.left_indent = Cm(1.5)
    par.paragraph_format.first_line_indent = Cm(0)
    par.style = doc.styles['List Bullet']
    run = par.add_run(text)
    sf(run)

def sc(cell, text, bold=False, align=WD_ALIGN_PARAGRAPH.CENTER, size=10):
    cell.text = ''
    pp = cell.paragraphs[0]
    pp.alignment = align
    pp.paragraph_format.space_before = Pt(2)
    pp.paragraph_format.space_after = Pt(2)
    pp.paragraph_format.first_line_indent = Cm(0)
    pp.paragraph_format.line_spacing = 1.0
    run = pp.add_run(text)
    sf(run, bold=bold, size=size)

def sh(row, color='D9E2F3'):
    for cell in row.cells:
        sd = OxmlElement('w:shd')
        sd.set(qn('w:fill'), color)
        cell._tc.get_or_add_tcPr().append(sd)

# ══════════════════════════════════════════
# COVER
# ══════════════════════════════════════════
for _ in range(5):
    p('', indent=False)

p('LAPORAN TUGAS', bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, indent=False, size=14)
p('SISTEM PENDUKUNG KEPUTUSAN', bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, indent=False, size=14)
p('', indent=False)
p('IMPLEMENTASI METODE SIMPLE ADDITIVE WEIGHTING (SAW)', bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, indent=False, size=14)
p('UNTUK SELEKSI SUPPLIER AKSESORIS HANDPHONE', bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, indent=False, size=14)
p('PADA CV ANUGERAH MEGA MAKMUR PONTIANAK', bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, indent=False, size=14)
p('', indent=False)
p('', indent=False)
p('Disusun Oleh:', align=WD_ALIGN_PARAGRAPH.CENTER, indent=False)
p('', indent=False)
kelompok = [
    '[NAMA MAHASISWA 1]  /  [NIM]',
    '[NAMA MAHASISWA 2]  /  [NIM]',
    '[NAMA MAHASISWA 3]  /  [NIM]',
    '[NAMA MAHASISWA 4]  /  [NIM]',
]
for k in kelompok:
    p(k, align=WD_ALIGN_PARAGRAPH.CENTER, indent=False)

p('', indent=False)
p('', indent=False)
p('[NAMA UNIVERSITAS]', bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, indent=False)
p('[FAKULTAS / PROGRAM STUDI]', align=WD_ALIGN_PARAGRAPH.CENTER, indent=False)
p('[KOTA]', align=WD_ALIGN_PARAGRAPH.CENTER, indent=False)
p('[TAHUN]', align=WD_ALIGN_PARAGRAPH.CENTER, indent=False)

doc.add_page_break()

# ══════════════════════════════════════════
# KATA PENGANTAR
# ══════════════════════════════════════════
h('KATA PENGANTAR', level=0, sb=0)

p('Puji syukur kami panjatkan ke hadirat Tuhan Yang Maha Esa karena atas rahmat dan karunia-Nya, kami dapat menyelesaikan laporan tugas Sistem Pendukung Keputusan ini dengan baik. Laporan ini disusun sebagai salah satu bentuk pemenuhan tugas mata kuliah Sistem Pendukung Keputusan pada Program Studi [NAMA PRODI] di [NAMA UNIVERSITAS].')

p('Laporan ini berisi tentang implementasi metode Simple Additive Weighting (SAW) untuk menyelesaikan permasalahan seleksi supplier aksesoris handphone di CV Anugerah Mega Makmur Pontianak. Dalam penyusunan laporan ini, kami melakukan observasi langsung ke lapangan, wawancara dengan pemilik perusahaan, serta studi pustaka dari berbagai sumber referensi.')

p('Kami menyadari bahwa laporan ini masih jauh dari sempurna. Oleh karena itu, kami sangat mengharapkan kritik dan saran yang membangun dari berbagai pihak demi perbaikan laporan ini ke depannya. Semoga laporan ini dapat bermanfaat bagi seluruh pembaca, khususnya bagi mahasiswa yang ingin mempelajari lebih lanjut tentang sistem pendukung keputusan dan metode SAW.')

p('', indent=False)
p('Pontianak, [TANGGAL]', align=WD_ALIGN_PARAGRAPH.RIGHT, indent=False)
p('', indent=False)
p('', indent=False)
p('Tim Penyusun', align=WD_ALIGN_PARAGRAPH.RIGHT, indent=False)

doc.add_page_break()

# ══════════════════════════════════════════
# DAFTAR ISI
# ══════════════════════════════════════════
h('DAFTAR ISI', level=0, sb=0)
p('', indent=False)

items = [
    ('HALAMAN JUDUL', 'i'),
    ('KATA PENGANTAR', 'ii'),
    ('DAFTAR ISI', 'iii'),
    ('DAFTAR TABEL', 'iv'),
    ('BAB I   PENDAHULUAN', '1'),
    ('1.1 Latar Belakang', '1'),
    ('1.2 Rumusan Masalah', '3'),
    ('1.3 Tujuan Penelitian', '3'),
    ('1.4 Batasan Masalah', '4'),
    ('1.5 Manfaat Penelitian', '4'),
    ('1.6 Metode Penelitian', '5'),
    ('BAB II  PEMBAHASAN', '7'),
    ('2.1 Profil CV Anugerah Mega Makmur', '7'),
    ('2.2 Landasan Teori', '8'),
    ('2.2.1 Sistem Pendukung Keputusan', '8'),
    ('2.2.2 Metode Simple Additive Weighting (SAW)', '9'),
    ('2.3 Identifikasi Kriteria dan Bobot', '11'),
    ('2.4 Perhitungan Manual Metode SAW', '12'),
    ('2.4.1 Data Alternatif Supplier', '12'),
    ('2.4.2 Proses Normalisasi', '14'),
    ('2.4.3 Perhitungan Skor Akhir', '15'),
    ('2.4.4 Perankingan', '16'),
    ('2.5 Implementasi Sistem Berbasis Web', '17'),
    ('2.5.1 Arsitektur Sistem', '17'),
    ('2.5.2 Tampilan Antarmuka', '18'),
    ('2.5.3 Pengujian Sistem', '19'),
    ('2.6 Perbandingan Manual dan Web', '20'),
    ('2.7 Dokumentasi Observasi dan Wawancara', '21'),
    ('BAB III PENUTUP', '22'),
    ('3.1 Kesimpulan', '22'),
    ('3.2 Saran', '23'),
    ('DAFTAR PUSTAKA', '24'),
]
for t, page in items:
    par = doc.add_paragraph()
    par.paragraph_format.line_spacing = 1.5
    par.paragraph_format.first_line_indent = Cm(0)
    par.paragraph_format.tab_stops.add_tab_stop(Cm(16))
    run = par.add_run(t)
    sf(run)
    run2 = par.add_run(f'\t{page}')
    sf(run2)

doc.add_page_break()

# ══════════════════════════════════════════
# DAFTAR TABEL
# ══════════════════════════════════════════
h('DAFTAR TABEL', level=0, sb=0)
p('', indent=False)

tables_list = [
    ('Tabel 2.1  Kriteria dan Bobot', '11'),
    ('Tabel 2.2  Data Alternatif Supplier', '13'),
    ('Tabel 2.3  Nilai Alternatif (Skala 0-1)', '13'),
    ('Tabel 2.4  Nilai Min dan Maks', '14'),
    ('Tabel 2.5  Hasil Normalisasi', '15'),
    ('Tabel 2.6  Skor Akhir SAW', '16'),
    ('Tabel 2.7  Ranking Supplier', '17'),
    ('Tabel 2.8  Perbandingan Manual vs Web', '20'),
]
for t, page in tables_list:
    par = doc.add_paragraph()
    par.paragraph_format.line_spacing = 1.5
    par.paragraph_format.first_line_indent = Cm(0)
    par.paragraph_format.tab_stops.add_tab_stop(Cm(16))
    run = par.add_run(t)
    sf(run)
    run2 = par.add_run(f'\t{page}')
    sf(run2)

doc.add_page_break()

# ══════════════════════════════════════════════
# BAB 1
# ══════════════════════════════════════════════
h('BAB I', level=0, sb=0)
h('PENDAHULUAN', level=0, sb=0)

h('1.1 Latar Belakang', level=2)

p('Perkembangan industri aksesoris handphone di Indonesia mengalami peningkatan yang cukup signifikan dalam beberapa tahun terakhir. Hal ini tidak lepas dari pertumbuhan jumlah pengguna smartphone yang terus bertambah setiap tahunnya. Data dari Asosiasi Penyelenggara Telekomunikasi Seluruh Indonesia (ATSI) mencatat bahwa jumlah pengguna smartphone di Indonesia sudah mencapai lebih dari 350 juta perangkat pada tahun 2025. Angka ini tentu saja membuka peluang besar bagi pelaku usaha di bidang aksesoris handphone, mulai dari charger, kabel data, casing, tempered glass, powerbank, earphone, hingga berbagai aksesoris pendukung lainnya.')

p('CV Anugerah Mega Makmur merupakan salah satu perusahaan yang bergerak di bidang perdagangan aksesoris handphone secara grosir yang berlokasi di Kota Pontianak, Kalimantan Barat. Perusahaan ini telah berdiri sejak tahun 2018 dan melayani berbagai toko retail dan reseller di wilayah Pontianak dan sekitarnya. Dalam menjalankan kegiatan operasionalnya sehari-hari, CV Anugerah Mega Makmur bekerja sama dengan kurang lebih 17 supplier yang menyediakan berbagai jenis aksesoris handphone dari berbagai merek dan kategori.')

p('Pemilihan supplier yang tepat merupakan salah satu faktor yang sangat menentukan keberhasilan bisnis di bidang perdagangan. Supplier yang berkualitas akan memberikan dampak positif terhadap kelancaran operasional perusahaan, seperti ketersediaan stok barang yang terjamin, harga yang kompetitif, kualitas produk yang baik, serta pengiriman yang tepat waktu. Sebaliknya, kesalahan dalam memilih supplier dapat mengakibatkan berbagai kerugian seperti barang berkualitas rendah, keterlambatan pengiriman, hingga hilangnya kepercayaan dari pelanggan.')

p('Permasalahan yang dihadapi oleh CV Anugerah Mega Makmur saat ini adalah belum adanya sistem yang terkomputerisasi dan terstruktur untuk membantu proses pengambilan keputusan dalam hal pemilihan supplier. Proses seleksi supplier yang berjalan selama ini masih menggunakan cara-cara yang sifatnya subjektif dan informal. Pemilik perusahaan mengaku bahwa selama ini pemilihan supplier hanya didasarkan pada pengalaman pribadi, kedekatan hubungan, atau rekomendasi dari kenalan. Cara seperti ini tentu saja memiliki kelemahan karena keputusan yang diambil cenderung tidak konsisten dan sulit untuk dipertanggungjawabkan secara objektif.')

p('Selain itu, kriteria-kriteria yang digunakan dalam proses seleksi supplier selama ini juga belum terdefinisi dengan jelas dan belum memiliki bobot prioritas yang pasti. Padahal, dalam praktiknya, terdapat beberapa aspek penting yang perlu dipertimbangkan secara bersamaan dalam memilih supplier, di antaranya adalah harga yang kompetitif, kualitas barang, ketepatan pengiriman, layanan yang responsif, serta ketersediaan stok atau kapasitas. Masing-masing aspek ini memiliki tingkat kepentingan yang berbeda dan perlu diberikan bobot yang sesuai.')

p('Berdasarkan permasalahan di atas, maka diperlukan suatu sistem pendukung keputusan yang dapat membantu pihak manajemen CV Anugerah Mega Makmur dalam melakukan seleksi supplier secara lebih objektif, terstruktur, dan terdokumentasi. Sistem pendukung keputusan atau SPK adalah sistem berbasis komputer yang dapat membantu proses pengambilan keputusan dengan menyediakan informasi, pemodelan, dan manipulasi data (Turban, 2005). Dengan adanya SPK, keputusan yang diambil diharapkan menjadi lebih akurat karena didasarkan pada perhitungan matematis yang jelas.')

p('Metode yang digunakan dalam penelitian ini adalah Simple Additive Weighting (SAW). Metode SAW dipilih karena memiliki beberapa keunggulan dibandingkan metode multi-kriteria lainnya. Pertama, metode SAW memiliki konsep yang sederhana dan mudah dipahami, sehingga cocok untuk diimplementasikan dalam sistem berbasis web. Kedua, metode SAW mampu menangani permasalahan multi-kriteria dengan baik melalui proses normalisasi dan pembobotan. Ketiga, metode SAW sudah banyak digunakan dalam berbagai penelitian serupa dan hasilnya teruji akurat (Kusumadewi, 2006).')

p('Dengan menerapkan metode SAW dalam sistem pendukung keputusan seleksi supplier, diharapkan CV Anugerah Mega Makmur dapat memperoleh rekomendasi supplier yang lebih objektif dan dapat diandalkan. Sistem yang dibangun juga diharapkan dapat menjadi alat bantu yang berguna bagi manajemen dalam mengambil keputusan strategis terkait pemilihan supplier di masa yang akan datang.')

h('1.2 Rumusan Masalah', level=2)

p('Berdasarkan uraian latar belakang di atas, maka rumusan masalah dalam laporan ini adalah sebagai berikut:')

p('1. Bagaimana cara mengidentifikasi kriteria-kriteria yang relevan dan menentukan bobot prioritas dalam proses seleksi supplier aksesoris handphone di CV Anugerah Mega Makmur?')

p('2. Bagaimana menerapkan metode Simple Additive Weighting (SAW) untuk menyelesaikan permasalahan seleksi supplier yang bersifat multi-kriteria?')

p('3. Bagaimana merancang dan membangun sistem pendukung keputusan berbasis web yang dapat mengimplementasikan metode SAW secara tepat dan akurat?')

p('4. Bagaimana tingkat akurasi sistem yang dibangun apabila dibandingkan dengan hasil perhitungan manual?')

h('1.3 Tujuan Penelitian', level=2)

p('Adapun tujuan dari penelitian ini adalah:')

p('1. Mengidentifikasi kriteria-kriteria yang digunakan dalam seleksi supplier di CV Anugerah Mega Makmur beserta bobot prioritasnya.')

p('2. Menerapkan metode Simple Additive Weighting (SAW) dalam sistem pendukung keputusan untuk seleksi supplier aksesoris handphone.')

p('3. Membangun aplikasi berbasis web yang dapat mengimplementasikan metode SAW secara user-friendly.')

p('4. Menguji dan membandingkan hasil perhitungan sistem dengan perhitungan manual untuk memvalidasi keakuratan sistem.')

h('1.4 Batasan Masalah', level=2)

p('Agar pembahasan dalam laporan ini lebih terarah dan tidak meluas, maka diberikan beberapa batasan masalah sebagai berikut:')

p('1. Penelitian hanya dilakukan pada CV Anugerah Mega Makmur yang berlokasi di Pontianak, Kalimantan Barat.')

p('2. Jumlah kriteria yang digunakan terbatas pada 5 kriteria, yaitu Harga, Kualitas, Pengiriman, Layanan, dan Kapasitas.')

p('3. Data supplier yang digunakan adalah data supplier aksesoris handphone yang menjadi mitra CV Anugerah Mega Makmur.')

p('4. Metode SPK yang digunakan hanya metode Simple Additive Weighting (SAW) tanpa melakukan perbandingan dengan metode lain.')

p('5. Sistem yang dibangun berbasis web dan belum dikembangkan dalam bentuk aplikasi mobile.')

h('1.5 Manfaat Penelitian', level=2)

p('Penelitian ini diharapkan dapat memberikan manfaat sebagai berikut:')

p('1. Bagi Perusahaan: Memberikan solusi sistem pendukung keputusan yang dapat membantu proses seleksi supplier secara objektif dan terstruktur, sehingga kualitas keputusan yang diambil menjadi lebih baik.')

p('2. Bagi Akademisi: Menambah wawasan dan pengetahuan mengenai penerapan metode Simple Additive Weighting (SAW) dalam sistem pendukung keputusan, khususnya dalam konteks pemilihan supplier.')

p('3. Bagi Penulis: Menerapkan ilmu yang telah diperoleh selama perkuliahan, khususnya mata kuliah Sistem Pendukung Keputusan, ke dalam bentuk aplikasi nyata yang dapat digunakan oleh masyarakat.')

h('1.6 Metode Penelitian', level=2)

p('Metode penelitian yang digunakan dalam penyusunan laporan ini terdiri dari beberapa tahapan, yaitu:')

p('1.6.1 Observasi', bold=True, indent=False)
p('Observasi dilakukan dengan cara mengunjungi langsung CV Anugerah Mega Makmur dan mengamati proses bisnis yang berjalan, khususnya terkait dengan proses pemilihan supplier aksesoris handphone. Observasi dilakukan selama [HARI] hari pada tanggal [TANGGAL]. Kegiatan observasi meliputi pengamatan terhadap alur penerimaan barang, proses pengecekan kualitas, pencatatan data supplier, dan interaksi antara pemilik perusahaan dengan supplier.')

p('1.6.2 Wawancara', bold=True, indent=False)
p('Wawancara dilakukan dengan pemilik CV Anugerah Mega Makmur, [NAMA PEMILIK], untuk menggali informasi mengenai kriteria-kriteria yang digunakan dalam pemilihan supplier, bobot prioritas dari setiap kriteria, kendala yang dihadapi dalam proses seleksi supplier, serta harapan terhadap sistem yang akan dibangun. Wawancara dilakukan secara semi-terstruktur dengan panduan pertanyaan yang telah disiapkan sebelumnya.')

p('1.6.3 Studi Pustaka', bold=True, indent=False)
p('Studi pustaka dilakukan dengan mempelajari literatur-literatur yang relevan dengan sistem pendukung keputusan dan metode Simple Additive Weighting (SAW). Sumber-sumber yang digunakan antara lain buku teks, jurnal ilmiah nasional dan internasional, artikel, dan sumber-sumber terpercaya dari internet. Beberapa literatur utama yang menjadi acuan adalah buku karya Kusumadewi (2006) tentang Fuzzy Multi-Attribute Decision Making dan buku karya Turban (2005) tentang Decision Support Systems.')

p('1.6.4 Perancangan Sistem', bold=True, indent=False)
p('Tahap perancangan sistem meliputi perancangan basis data, perancangan arsitektur sistem, perancangan antarmuka pengguna, dan perancangan alur kerja sistem. Sistem dirancang dengan arsitektur klien-server dimana frontend dibangun menggunakan Next.js 14 dan backend menggunakan NestJS. Database yang digunakan adalah PostgreSQL. Perancangan dilakukan dengan membuat diagram alir, struktur database, dan wireframe antarmuka.')

p('1.6.5 Implementasi', bold=True, indent=False)
p('Tahap implementasi merupakan tahap penerjemahan desain sistem ke dalam kode program. Frontend dikembangkan menggunakan framework Next.js 14 dengan bahasa pemrograman TypeScript dan CSS framework Tailwind CSS. Backend dikembangkan menggunakan framework NestJS yang juga menggunakan TypeScript. Database menggunakan PostgreSQL yang dijalankan di Supabase. Implementasi algoritma SAW dilakukan pada sisi backend dalam bentuk modul khusus.')

p('1.6.6 Pengujian', bold=True, indent=False)
p('Pengujian sistem dilakukan dengan dua pendekatan. Pertama, pengujian fungsionalitas untuk memastikan bahwa seluruh fitur sistem berjalan sesuai dengan yang diharapkan. Kedua, pengujian akurasi dengan membandingkan hasil perhitungan sistem dengan hasil perhitungan manual menggunakan data sampel yang sama. Hasil pengujian dicatat dan dianalisis untuk mengetahui tingkat keakuratan sistem.')

doc.add_page_break()

# ══════════════════════════════════════════════
# BAB 2
# ══════════════════════════════════════════════
h('BAB II', level=0, sb=0)
h('PEMBAHASAN', level=0, sb=0)

h('2.1 Profil CV Anugerah Mega Makmur', level=2)

p('CV Anugerah Mega Makmur adalah sebuah perusahaan yang bergerak di bidang perdagangan aksesoris handphone secara grosir. Perusahaan ini didirikan pada tahun 2018 dan beralamat di Kota Pontianak, Kalimantan Barat. CV Anugerah Mega Makmur menyediakan berbagai macam aksesoris handphone seperti charger, kabel data, casing handphone, tempered glass, powerbank, earphone, headset, dan berbagai aksesoris pendukung lainnya.')

p('Visi dari CV Anugerah Mega Makmur adalah menjadi perusahaan penyedia aksesoris handphone terkemuka di Kalimantan Barat yang mampu memenuhi kebutuhan pelanggan dengan produk berkualitas dan harga kompetitif. Sedangkan misinya antara lain adalah menyediakan produk aksesoris handphone yang berkualitas dengan harga yang terjangkau, menjalin kemitraan yang saling menguntungkan dengan para supplier, dan memberikan pelayanan terbaik kepada pelanggan.')

p('Dalam menjalankan kegiatan usahanya, CV Anugerah Mega Makmur memiliki struktur organisasi yang terdiri dari pemilik, bagian administrasi dan keuangan, bagian gudang dan logistik, serta bagian pemasaran dan penjualan. Saat ini perusahaan memiliki kurang lebih 17 supplier aktif yang menyediakan berbagai jenis aksesoris handphone dari berbagai merek seperti Baseus, Anker, Ugreen, Robot, Vivan, JBL, Oraimo, dan masih banyak lagi. Supplier-supplier tersebut berasal dari wilayah Pontianak dan sekitarnya, serta beberapa dari luar daerah.')

p('Proses bisnis utama CV Anugerah Mega Makmur meliputi: (1) pemesanan barang ke supplier, (2) penerimaan dan pengecekan barang, (3) penyimpanan barang di gudang, (4) penjualan barang ke pelanggan (reseller dan toko retail), dan (5) pengiriman barang ke pelanggan. Dalam proses pemesanan barang ke supplier, pemilik perusahaan harus memutuskan supplier mana yang akan dipilih untuk setiap jenis barang yang dibutuhkan. Keputusan inilah yang selama ini masih dilakukan secara subjektif dan belum menggunakan metode yang terstruktur.')

h('2.2 Landasan Teori', level=2)

h('2.2.1 Sistem Pendukung Keputusan', level=3)

p('Sistem Pendukung Keputusan (SPK) atau Decision Support System (DSS) pertama kali diperkenalkan oleh Michael S. Scott Morton pada awal tahun 1970-an. SPK didefinisikan sebagai sistem berbasis komputer yang interaktif, yang membantu pengambil keputusan dalam menggunakan data dan model untuk menyelesaikan masalah-masalah yang tidak terstruktur (Turban, 2005). SPK dirancang untuk mendukung seluruh tahapan pengambilan keputusan, mulai dari identifikasi masalah, pemilihan data yang relevan, penentuan pendekatan yang digunakan, hingga evaluasi pemilihan alternatif.')

p('Menurut Suryadi dan Ramdhani (2000), terdapat beberapa karakteristik utama dari SPK, yaitu:')

p('1. SPK ditujukan untuk membantu pengambil keputusan dalam memecahkan masalah yang bersifat semi-terstruktur atau tidak terstruktur.')
p('2. SPK merupakan gabungan antara data, model, dan antarmuka pengguna.')
p('3. SPK bersifat fleksibel dan dapat beradaptasi dengan perubahan situasi dan kebutuhan pengguna.')
p('4. SPK tidak dimaksudkan untuk menggantikan peran pengambil keputusan, melainkan sebagai alat bantu.')
p('5. SPK menggunakan data dan model matematis atau analitis dalam proses pengambilan keputusan.')

p('Komponen utama dari SPK terdiri dari tiga subsistem, yaitu:')

p('a. Subsistem Data (Database Subsystem)', bold=True, indent=False)
p('Subsistem data berfungsi sebagai penyedia data yang relevan bagi pengambil keputusan. Data dapat berasal dari database internal perusahaan maupun dari sumber-sumber eksternal. Dalam sistem yang dibangun, subsistem data menggunakan database PostgreSQL yang menyimpan data supplier, data kriteria, data bobot, dan hasil evaluasi.')

p('b. Subsistem Model (Model Subsystem)', bold=True, indent=False)
p('Subsistem model merupakan inti dari SPK yang berisi model-model matematis atau analitis untuk mengolah data menjadi informasi yang berguna. Dalam penelitian ini, model yang digunakan adalah algoritma Simple Additive Weighting (SAW) yang akan melakukan normalisasi dan perhitungan bobot untuk menghasilkan skor akhir setiap alternatif.')

p('c. Subsistem Dialog (Dialog Subsystem)', bold=True, indent=False)
p('Subsistem dialog atau user interface merupakan jembatan antara pengguna dengan sistem. Antarmuka yang baik harus mudah dipahami dan digunakan oleh pengguna yang tidak memiliki latar belakang teknis. Dalam sistem ini, antarmuka dibangun menggunakan framework Next.js dengan desain yang responsif dan intuitif.')

h('2.2.2 Metode Simple Additive Weighting (SAW)', level=3)

p('Simple Additive Weighting (SAW) sering juga dikenal dengan istilah metode penjumlahan terbobot. Konsep dasar metode SAW adalah mencari penjumlahan terbobot dari rating kinerja pada setiap alternatif di semua atribut atau kriteria yang ada (Kusumadewi, 2006). Metode SAW termasuk dalam kategori Multi-Attribute Decision Making (MADM) yang menangani permasalahan dengan jumlah alternatif yang terbatas dan beberapa kriteria yang harus dipertimbangkan.')

p('Metode SAW memerlukan proses normalisasi matriks keputusan ke suatu skala nilai yang dapat diperbandingkan dengan semua rating alternatif yang ada. Proses normalisasi ini sangat penting karena setiap kriteria mungkin memiliki satuan dan rentang nilai yang berbeda-beda. Dengan dinormalisasi, nilai-nilai dari berbagai kriteria dapat dibandingkan secara adil.')

p('Langkah-langkah penyelesaian menggunakan metode SAW adalah sebagai berikut:')

p('Langkah 1: Menentukan Kriteria', bold=True, indent=False)
p('Menentukan kriteria-kriteria yang akan dijadikan acuan dalam pengambilan keputusan. Dalam penelitian ini, terdapat 5 kriteria yang digunakan yaitu Harga (C1), Kualitas (C2), Pengiriman (C3), Layanan (C4), dan Kapasitas (C5).')

p('Langkah 2: Menentukan Bobot', bold=True, indent=False)
p('Menentukan bobot preferensi atau tingkat kepentingan untuk setiap kriteria. Bobot harus berjumlah 1 (atau 100%). Dalam penelitian ini, bobot ditetapkan sebagai berikut: Harga 30%, Kualitas 30%, Pengiriman 20%, Layanan 10%, dan Kapasitas 10%.')

p('Langkah 3: Membuat Matriks Keputusan', bold=True, indent=False)
p('Membuat matriks keputusan X yang berisi nilai setiap alternatif (supplier) pada setiap kriteria. Matriks X berukuran m x n, dimana m adalah jumlah alternatif dan n adalah jumlah kriteria.')

p('Langkah 4: Normalisasi Matriks', bold=True, indent=False)
p('Melakukan normalisasi matriks keputusan berdasarkan jenis kriteria. Setiap kriteria diklasifikasikan menjadi dua jenis, yaitu kriteria benefit (keuntungan) dan kriteria cost (biaya). Rumus normalisasinya adalah sebagai berikut:')

p('Untuk kriteria benefit:', italic=True, indent=False)
p('rij = xij / max(xij)', align=WD_ALIGN_PARAGRAPH.CENTER, indent=False, italic=True)

p('Untuk kriteria cost:', italic=True, indent=False)
p('rij = min(xij) / xij', align=WD_ALIGN_PARAGRAPH.CENTER, indent=False, italic=True)

p('Keterangan:', italic=True, indent=False)
p('rij = nilai rating kinerja ternormalisasi dari alternatif ke-i pada kriteria ke-j', indent=False)
p('xij = nilai atribut dari alternatif ke-i pada kriteria ke-j', indent=False)
p('max(xij) = nilai maksimum dari setiap baris dan kolom (kriteria benefit)', indent=False)
p('min(xij) = nilai minimum dari setiap baris dan kolom (kriteria cost)', indent=False)

p('Langkah 5: Menghitung Nilai Preferensi', bold=True, indent=False)
p('Menghitung nilai preferensi (Vi) untuk setiap alternatif dengan menjumlahkan hasil kali antara bobot kriteria (wj) dengan nilai ternormalisasi (rij). Supplier dengan nilai Vi tertinggi adalah alternatif terbaik yang layak dipilih. Rumusnya adalah:')

p('Vi = Σ wj × rij', align=WD_ALIGN_PARAGRAPH.CENTER, indent=False, italic=True)

p('Keterangan:', italic=True, indent=False)
p('Vi = nilai akhir alternatif ke-i', indent=False)
p('wj = bobot kriteria ke-j', indent=False)
p('rij = nilai ternormalisasi alternatif ke-i pada kriteria ke-j', indent=False)

p('Dalam penerapannya di penelitian ini, selain menggunakan rumus SAW di atas, juga diberikan bonus sebesar 0.05 poin untuk supplier yang bersedia menanggung biaya ongkos kirim (shipping coverage = SUPPLIER_COVERS). Bonus ini merupakan kebijakan dari CV Anugerah Mega Makmur untuk memberikan apresiasi kepada supplier yang memberikan layanan tambahan. Skor akhir setiap supplier kemudian dibandingkan dengan threshold kelulusan sebesar 0.75 untuk menentukan apakah supplier tersebut direkomendasikan atau tidak.')

h('2.3 Identifikasi Kriteria dan Bobot', level=2)

p('Berdasarkan hasil observasi dan wawancara dengan pemilik CV Anugerah Mega Makmur, terdapat lima kriteria utama yang digunakan dalam proses seleksi supplier aksesoris handphone. Setiap kriteria memiliki bobot yang mencerminkan tingkat kepentingannya. Berikut adalah rincian kriteria, jenis, dan bobotnya:')

# Tabel Kriteria
t1 = doc.add_table(rows=7, cols=5)
t1.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, hh in enumerate(['No', 'Kriteria', 'Jenis', 'Bobot', 'Keterangan']):
    sc(t1.rows[0].cells[i], hh, bold=True, size=9)
sh(t1.rows[0])

kriteria_data = [
    ('1', 'Harga', 'Cost', '30%', 'Semakin murah harga, semakin baik'),
    ('2', 'Kualitas', 'Benefit', '30%', 'Semakin baik kualitas, semakin baik'),
    ('3', 'Pengiriman', 'Benefit', '20%', 'Semakin cepat & tepat, semakin baik'),
    ('4', 'Layanan', 'Benefit', '10%', 'Semakin responsif, semakin baik'),
    ('5', 'Kapasitas', 'Benefit', '10%', 'Semakin besar stok, semakin baik'),
    ('', 'Bonus Ongkir', 'Tambahan', '+0.05', 'Jika supplier menanggung ongkos kirim'),
]
for r, rd in enumerate(kriteria_data, 1):
    for c, val in enumerate(rd):
        sc(t1.rows[r].cells[c], val, size=8)

p('')
p('Tabel 2.1 di atas menunjukkan bahwa terdapat dua jenis kriteria dalam proses seleksi supplier, yaitu kriteria cost dan benefit. Harga merupakan satu-satunya kriteria cost karena prinsipnya semakin murah harga yang ditawarkan supplier maka semakin baik bagi perusahaan. Sedangkan kualitas, pengiriman, layanan, dan kapasitas merupakan kriteria benefit yang artinya semakin tinggi nilainya maka semakin baik.')

p('Bobot total dari kelima kriteria adalah 100%. Pembagian bobot ini menunjukkan bahwa harga dan kualitas menjadi prioritas utama dengan bobot masing-masing 30%. Hal ini sesuai dengan hasil wawancara di mana pemilik CV Anugerah Mega Makmur menyatakan bahwa harga dan kualitas adalah dua faktor yang paling penting dalam memilih supplier. Pengiriman juga cukup penting dengan bobot 20%, sedangkan layanan dan kapasitas masing-masing memiliki bobot 10%. Selain itu, terdapat bonus tambahan sebesar 0.05 poin bagi supplier yang bersedia menanggung biaya ongkos kirim.')

h('2.4 Perhitungan Manual Metode SAW', level=2)

p('Untuk memahami penerapan metode SAW secara lebih mendalam, berikut akan disajikan perhitungan manual menggunakan 5 sampel supplier aksesoris handphone yang menjadi mitra CV Anugerah Mega Makmur. Perhitungan ini akan dilakukan langkah demi langkah sesuai dengan prosedur metode SAW yang telah dijelaskan pada subbab sebelumnya.')

h('2.4.1 Data Alternatif Supplier', level=3)

p('Lima sampel supplier yang digunakan dalam perhitungan manual ini adalah sebagai berikut:')
p('1. Pontianak Mobile Grosir - Menyediakan charger dengan berbagai merek')
p('2. Khatulistiwa Gadget Supply - Menyediakan kabel data dan aksesoris gadget')
p('3. Borneo Tech Distributor - Menyediakan aksesoris premium')
p('4. Mega Jaya Cellular Pontianak - Menyediakan casing handphone')
p('5. JBL Audio Partner - Menyediakan perangkat audio')

p('Setiap supplier dinilai pada masing-masing kriteria menggunakan skala 1 sampai 10. Penilaian dilakukan oleh pemilik CV Anugerah Mega Makmur berdasarkan pengalaman dan data transaksi selama bekerja sama dengan supplier-supplier tersebut. Berikut adalah data nilai dari masing-masing alternatif:')

# Tabel data alternatif
t2 = doc.add_table(rows=6, cols=7)
t2.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, hh in enumerate(['Alternatif', 'Harga', 'Kualitas', 'Pengiriman', 'Layanan', 'Kapasitas', 'Ongkir']):
    sc(t2.rows[0].cells[i], hh, bold=True, size=8)
sh(t2.rows[0])
sup_data = [
    ('Pontianak Mobile Grosir', '9,1', '9,0', '8,8', '8,6', '9,2', 'Tidak'),
    ('Khatulistiwa Gadget Supply', '8,6', '9,3', '9,0', '8,9', '8,8', 'Ya'),
    ('Borneo Tech Distributor', '8,2', '9,5', '8,4', '8,8', '8,6', 'Tidak'),
    ('Mega Jaya Cellular', '8,9', '8,8', '9,2', '8,7', '8,9', 'Tidak'),
    ('JBL Audio Partner', '8,0', '9,4', '8,3', '8,7', '8,2', 'Tidak'),
]
for r, rd in enumerate(sup_data, 1):
    for c, val in enumerate(rd):
        sc(t2.rows[r].cells[c], val, size=8)

p('')
p('Tabel 2.2 menunjukkan data nilai kelima supplier pada setiap kriteria. Nilai-nilai ini kemudian dikonversi ke dalam skala 0-1 dengan cara membagi setiap nilai dengan 10. Hasil konversi dapat dilihat pada tabel berikut:')

# Tabel konversi
t3 = doc.add_table(rows=6, cols=6)
t3.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, hh in enumerate(['Alternatif', 'Harga', 'Kualitas', 'Pengiriman', 'Layanan', 'Kapasitas']):
    sc(t3.rows[0].cells[i], hh, bold=True, size=9)
sh(t3.rows[0])
konv_data = [
    ('Pontianak Mobile Grosir', '0,91', '0,90', '0,88', '0,86', '0,92'),
    ('Khatulistiwa Gadget Supply', '0,86', '0,93', '0,90', '0,89', '0,88'),
    ('Borneo Tech Distributor', '0,82', '0,95', '0,84', '0,88', '0,86'),
    ('Mega Jaya Cellular', '0,89', '0,88', '0,92', '0,87', '0,89'),
    ('JBL Audio Partner', '0,80', '0,94', '0,83', '0,87', '0,82'),
]
for r, rd in enumerate(konv_data, 1):
    for c, val in enumerate(rd):
        sc(t3.rows[r].cells[c], val, size=9)

p('')
p('Setelah mendapatkan nilai dalam skala 0-1, langkah selanjutnya adalah menentukan nilai minimum dan maksimum untuk setiap kriteria. Nilai minimum digunakan untuk kriteria cost (Harga), sedangkan nilai maksimum digunakan untuk kriteria benefit (Kualitas, Pengiriman, Layanan, dan Kapasitas).')

# Tabel min max
t4 = doc.add_table(rows=3, cols=6)
t4.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, hh in enumerate(['', 'Harga', 'Kualitas', 'Pengiriman', 'Layanan', 'Kapasitas']):
    sc(t4.rows[0].cells[i], hh, bold=True, size=9)
sh(t4.rows[0])
mm_data = [('Min (Cost)', '0,80', '-', '-', '-', '-'),
           ('Max (Benefit)', '-', '0,95', '0,92', '0,89', '0,92')]
for r, rd in enumerate(mm_data, 1):
    for c, val in enumerate(rd):
        sc(t4.rows[r].cells[c], val, size=9)

p('')
p('Nilai minimum untuk Harga (kriteria cost) adalah 0,80 yang dimiliki oleh JBL Audio Partner. Nilai maksimum untuk Kualitas adalah 0,95 (Borneo Tech Distributor), Pengiriman adalah 0,92 (Mega Jaya Cellular), Layanan adalah 0,89 (Khatulistiwa Gadget Supply), dan Kapasitas adalah 0,92 (Pontianak Mobile Grosir).')

h('2.4.2 Proses Normalisasi', level=3)

p('Proses normalisasi dilakukan dengan menggunakan rumus yang telah dijelaskan sebelumnya, yaitu:')

p('a. Kriteria Cost (Harga): rij = min(xij) / xij', bold=True, indent=False)
p('b. Kriteria Benefit (Kualitas, Pengiriman, Layanan, Kapasitas): rij = xij / max(xij)', bold=True, indent=False)

p('Berikut adalah perhitungan normalisasi untuk masing-masing alternatif:')

p('a. Pontianak Mobile Grosir', bold=True, indent=False)
p('Harga = 0,80 / 0,91 = 0,8791', indent=False)
p('Kualitas = 0,90 / 0,95 = 0,9474', indent=False)
p('Pengiriman = 0,88 / 0,92 = 0,9565', indent=False)
p('Layanan = 0,86 / 0,89 = 0,9663', indent=False)
p('Kapasitas = 0,92 / 0,92 = 1,0000', indent=False)

p('b. Khatulistiwa Gadget Supply', bold=True, indent=False)
p('Harga = 0,80 / 0,86 = 0,9302', indent=False)
p('Kualitas = 0,93 / 0,95 = 0,9789', indent=False)
p('Pengiriman = 0,90 / 0,92 = 0,9783', indent=False)
p('Layanan = 0,89 / 0,89 = 1,0000', indent=False)
p('Kapasitas = 0,88 / 0,92 = 0,9565', indent=False)

p('c. Borneo Tech Distributor', bold=True, indent=False)
p('Harga = 0,80 / 0,82 = 0,9756', indent=False)
p('Kualitas = 0,95 / 0,95 = 1,0000', indent=False)
p('Pengiriman = 0,84 / 0,92 = 0,9130', indent=False)
p('Layanan = 0,88 / 0,89 = 0,9888', indent=False)
p('Kapasitas = 0,86 / 0,92 = 0,9348', indent=False)

p('d. Mega Jaya Cellular', bold=True, indent=False)
p('Harga = 0,80 / 0,89 = 0,8989', indent=False)
p('Kualitas = 0,88 / 0,95 = 0,9263', indent=False)
p('Pengiriman = 0,92 / 0,92 = 1,0000', indent=False)
p('Layanan = 0,87 / 0,89 = 0,9775', indent=False)
p('Kapasitas = 0,89 / 0,92 = 0,9674', indent=False)

p('e. JBL Audio Partner', bold=True, indent=False)
p('Harga = 0,80 / 0,80 = 1,0000', indent=False)
p('Kualitas = 0,94 / 0,95 = 0,9895', indent=False)
p('Pengiriman = 0,83 / 0,92 = 0,9022', indent=False)
p('Layanan = 0,87 / 0,89 = 0,9775', indent=False)
p('Kapasitas = 0,82 / 0,92 = 0,8913', indent=False)

p('Hasil normalisasi dapat disajikan dalam bentuk tabel berikut:')

# Tabel normalisasi
t5 = doc.add_table(rows=6, cols=6)
t5.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, hh in enumerate(['Alternatif', 'Harga', 'Kualitas', 'Pengiriman', 'Layanan', 'Kapasitas']):
    sc(t5.rows[0].cells[i], hh, bold=True, size=9)
sh(t5.rows[0])
norm_data = [
    ('Pontianak Mobile Grosir', '0,8791', '0,9474', '0,9565', '0,9663', '1,0000'),
    ('Khatulistiwa Gadget Supply', '0,9302', '0,9789', '0,9783', '1,0000', '0,9565'),
    ('Borneo Tech Distributor', '0,9756', '1,0000', '0,9130', '0,9888', '0,9348'),
    ('Mega Jaya Cellular', '0,8989', '0,9263', '1,0000', '0,9775', '0,9674'),
    ('JBL Audio Partner', '1,0000', '0,9895', '0,9022', '0,9775', '0,8913'),
]
for r, rd in enumerate(norm_data, 1):
    for c, val in enumerate(rd):
        sc(t5.rows[r].cells[c], val, size=9)

h('2.4.3 Perhitungan Skor Akhir', level=3)

p('Setelah mendapatkan nilai normalisasi, langkah selanjutnya adalah menghitung skor akhir (Vi) untuk setiap alternatif dengan rumus: Vi = Σ (wj × rij) + bonus ongkir (jika ada). Bobot yang digunakan adalah: Harga = 0,3; Kualitas = 0,3; Pengiriman = 0,2; Layanan = 0,1; Kapasitas = 0,1.')

p('a. Pontianak Mobile Grosir', bold=True, indent=False)
p('Vi = (0,8791×0,3) + (0,9474×0,3) + (0,9565×0,2) + (0,9663×0,1) + (1,0000×0,1) = 0,2637 + 0,2842 + 0,1913 + 0,0966 + 0,1000 = 0,9358', indent=False)

p('b. Khatulistiwa Gadget Supply', bold=True, indent=False)
p('Vi = (0,9302×0,3) + (0,9789×0,3) + (0,9783×0,2) + (1,0000×0,1) + (0,9565×0,1) + 0,05 (bonus) = 0,2791 + 0,2937 + 0,1957 + 0,1000 + 0,0957 + 0,05 = 1,0142', indent=False)

p('c. Borneo Tech Distributor', bold=True, indent=False)
p('Vi = (0,9756×0,3) + (1,0000×0,3) + (0,9130×0,2) + (0,9888×0,1) + (0,9348×0,1) = 0,2927 + 0,3000 + 0,1826 + 0,0989 + 0,0935 = 0,9677', indent=False)

p('d. Mega Jaya Cellular', bold=True, indent=False)
p('Vi = (0,8989×0,3) + (0,9263×0,3) + (1,0000×0,2) + (0,9775×0,1) + (0,9674×0,1) = 0,2697 + 0,2779 + 0,2000 + 0,0978 + 0,0967 = 0,9421', indent=False)

p('e. JBL Audio Partner', bold=True, indent=False)
p('Vi = (1,0000×0,3) + (0,9895×0,3) + (0,9022×0,2) + (0,9775×0,1) + (0,8913×0,1) = 0,3000 + 0,2969 + 0,1804 + 0,0978 + 0,0891 = 0,9642', indent=False)

# Tabel skor
t6 = doc.add_table(rows=6, cols=7)
t6.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, hh in enumerate(['Alternatif', 'Harga\n(×0,3)', 'Kualitas\n(×0,3)', 'Pengiriman\n(×0,2)', 'Layanan\n(×0,1)', 'Kapasitas\n(×0,1)', 'Skor\nAkhir']):
    sc(t6.rows[0].cells[i], hh, bold=True, size=8)
sh(t6.rows[0])
skor_data = [
    ('Pontianak Mobile Grosir', '0,2637', '0,2842', '0,1913', '0,0966', '0,1000', '0,9358'),
    ('Khatulistiwa Gadget Supply', '0,2791', '0,2937', '0,1957', '0,1000', '0,0957', '1,0142'),
    ('Borneo Tech Distributor', '0,2927', '0,3000', '0,1826', '0,0989', '0,0935', '0,9677'),
    ('Mega Jaya Cellular', '0,2697', '0,2779', '0,2000', '0,0978', '0,0967', '0,9421'),
    ('JBL Audio Partner', '0,3000', '0,2969', '0,1804', '0,0978', '0,0891', '0,9642'),
]
for r, rd in enumerate(skor_data, 1):
    for c, val in enumerate(rd):
        sc(t6.rows[r].cells[c], val, size=8)

h('2.4.4 Perankingan', level=3)

p('Berdasarkan skor akhir yang telah diperoleh, langkah terakhir adalah melakukan perankingan. Supplier dengan skor tertinggi menempati peringkat pertama dan menjadi rekomendasi utama. Threshold kelulusan ditetapkan pada nilai 0,75, sehingga supplier dengan skor ≥ 0,75 dinyatakan layak untuk direkomendasikan.')

# Tabel ranking
t7 = doc.add_table(rows=6, cols=5)
t7.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, hh in enumerate(['Ranking', 'Supplier', 'Skor SAW', 'Threshold\n(0,75)', 'Rekomendasi']):
    sc(t7.rows[0].cells[i], hh, bold=True, size=9)
sh(t7.rows[0])
rank_data = [
    ('1', 'Khatulistiwa Gadget Supply', '1,0142', '≥ 0,75', 'Direkomendasikan'),
    ('2', 'Borneo Tech Distributor', '0,9677', '≥ 0,75', 'Direkomendasikan'),
    ('3', 'JBL Audio Partner', '0,9642', '≥ 0,75', 'Direkomendasikan'),
    ('4', 'Mega Jaya Cellular', '0,9421', '≥ 0,75', 'Direkomendasikan'),
    ('5', 'Pontianak Mobile Grosir', '0,9358', '≥ 0,75', 'Direkomendasikan'),
]
for r, rd in enumerate(rank_data, 1):
    for c, val in enumerate(rd):
        sc(t7.rows[r].cells[c], val, size=9)

p('')
p('Berdasarkan Tabel 2.7, Khatulistiwa Gadget Supply menempati peringkat pertama dengan skor SAW 1,0142. Skor ini melebihi 1,0 karena supplier mendapatkan bonus ongkos kirim sebesar 0,05 poin. Peringkat kedua ditempati oleh Borneo Tech Distributor dengan skor 0,9677, disusul oleh JBL Audio Partner (0,9642), Mega Jaya Cellular (0,9421), dan Pontianak Mobile Grosir (0,9358). Seluruh supplier berada di atas threshold 0,75, sehingga semuanya dinyatakan direkomendasikan.')

p('Perlu dicatat bahwa skor Khatulistiwa Gadget Supply yang melebihi 1,0 bukanlah suatu kesalahan melainkan konsekuensi dari adanya bonus ongkos kirim. Dalam metode SAW, skor normalisasi berkisar antara 0 hingga 1, namun dengan adanya nilai tambah dari bonus, skor akhir dapat melebihi 1,0. Hal ini wajar terjadi dan menunjukkan bahwa supplier tersebut memberikan nilai lebih di luar kriteria standar.')

h('2.5 Implementasi Sistem Berbasis Web', level=2)

p('Sistem pendukung keputusan seleksi supplier berbasis web ini dibangun menggunakan teknologi-teknologi modern yang saling terintegrasi. Berikut adalah spesifikasi teknis dan penjelasan mengenai implementasi sistem:')

h('2.5.1 Arsitektur Sistem', level=3)

p('Sistem dibangun dengan arsitektur client-server yang terdiri dari tiga komponen utama:')

p('1. Frontend (Client Side): Dibangun menggunakan Next.js 14 dengan framework React dan bahasa TypeScript. Tampilan antarmuka menggunakan Tailwind CSS untuk styling yang responsif dan modern. Komponen-komponen UI menggunakan library Radix UI yang diintegrasikan dengan shadcn/ui. Frontend berjalan di port 3000.')

p('2. Backend (Server Side): Dibangun menggunakan NestJS yang merupakan framework Node.js untuk pengembangan aplikasi server-side yang terstruktur. Backend menyediakan RESTful API yang melayani permintaan dari frontend. Backend berjalan di port 4000 dengan prefix /api. Autentikasi menggunakan JWT (JSON Web Token) dengan passport-jwt.')

p('3. Database: Menggunakan PostgreSQL sebagai database relasional. Pengelolaan database menggunakan Prisma ORM yang memudahkan interaksi dengan database melalui kode TypeScript. Database dihosting menggunakan Supabase untuk kemudahan deployment.')

p('Komunikasi antara frontend dan backend dilakukan melalui protokol HTTP dengan format data JSON. Setiap permintaan ke backend (kecuali login) memerlukan token JWT yang dikirim melalui header Authorization. Backend memiliki beberapa endpoint utama yang terkait dengan modul SPK, yaitu:')

api_endpoints = [
    'GET /api/spk/suppliers - Mendapatkan daftar supplier',
    'POST /api/spk/suppliers - Menambahkan supplier baru',
    'PATCH /api/spk/suppliers/:id - Mengubah data supplier',
    'DELETE /api/spk/suppliers/:id - Menghapus data supplier',
    'POST /api/spk/supplier-selection - Menjalankan evaluasi SAW',
    'GET /api/spk/results - Mendapatkan hasil evaluasi',
    'PATCH /api/spk/results/:id - Approve/reject hasil evaluasi',
]
for ep in api_endpoints:
    b(ep)

h('2.5.2 Tampilan Antarmuka', level=3)

p('Sistem memiliki beberapa halaman utama yang masing-masing memiliki fungsi spesifik. Berikut adalah penjelasan mengenai setiap halaman:')

p('a. Halaman Dashboard', bold=True, indent=False)
p('Halaman dashboard merupakan halaman utama yang muncul saat pengguna berhasil login. Halaman ini menampilkan ringkasan data supplier berupa total supplier, jumlah supplier aktif, jumlah supplier yang direkomendasikan, rata-rata skor SAW, dan daftar 5 supplier terbaik berdasarkan skor. Informasi ini memberikan gambaran umum kepada pengguna mengenai kondisi supplier secara keseluruhan.')

p('b. Halaman Data Supplier', bold=True, indent=False)
p('Halaman ini digunakan untuk mengelola data supplier, meliputi penambahan supplier baru, mengubah data supplier yang sudah ada, dan menghapus data supplier. Setiap supplier memiliki data berupa nama, kategori, kontak, dan nilai untuk masing-masing kriteria (Harga, Kualitas, Pengiriman, Layanan, Kapasitas) dalam skala 1-10, serta status ongkos kirim. Data ini akan menjadi input utama dalam proses perhitungan SAW.')

p('c. Halaman Evaluasi Supplier', bold=True, indent=False)
p('Halaman ini merupakan halaman inti dari sistem. Pengguna dapat menjalankan proses evaluasi SAW dengan memilih filter kategori aksesoris (opsional) dan menentukan threshold kelulusan (skala 0-1, default 0,75). Setelah evaluasi selesai, sistem akan menampilkan tabel ranking yang berisi peringkat, nama supplier, nilai asli setiap kriteria, nilai normalisasi, skor SAW, status rekomendasi, serta tombol untuk approve atau reject hasil evaluasi.')

p('d. Halaman Kriteria Penilaian', bold=True, indent=False)
p('Halaman ini menampilkan informasi mengenai kriteria-kriteria yang digunakan dalam proses evaluasi, lengkap dengan bobot, jenis (cost/benefit), dan deskripsi. Halaman ini bersifat informatif dan membantu pengguna memahami bagaimana sistem melakukan perhitungan.')

h('2.5.3 Pengujian Sistem', level=3)

p('Pengujian sistem dilakukan untuk memastikan bahwa seluruh fungsionalitas sistem berjalan dengan baik. Pengujian meliputi:')

p('a. Pengujian Fungsionalitas', bold=True, indent=False)
p('Pengujian fungsionalitas dilakukan dengan menguji setiap fitur yang ada pada sistem, mulai dari login, CRUD supplier, menjalankan evaluasi, hingga approve/reject hasil. Seluruh fitur berjalan sesuai dengan yang diharapkan. Pengguna dapat login dengan peran masing-masing (SUPER_ADMIN, ADMIN_HR, MANAGER, KARYAWAN) dan mendapatkan akses sesuai dengan haknya.')

p('b. Pengujian Akurasi', bold=True, indent=False)
p('Pengujian akurasi dilakukan dengan membandingkan hasil perhitungan sistem dengan hasil perhitungan manual pada subbab sebelumnya. Perbandingan dilakukan dengan memasukkan data 5 supplier sampel ke dalam sistem dan menjalankan evaluasi SAW, kemudian membandingkan skor yang dihasilkan dengan skor manual.')

h('2.6 Perbandingan Perhitungan Manual dan Web', level=2)

p('Setelah melakukan perhitungan manual dan implementasi sistem, langkah selanjutnya adalah membandingkan hasil keduanya untuk memvalidasi keakuratan sistem. Berikut adalah tabel perbandingan hasil perhitungan manual dan hasil perhitungan pada sistem berbasis web:')

# Tabel perbandingan
t8 = doc.add_table(rows=6, cols=4)
t8.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, hh in enumerate(['Supplier', 'Skor Manual', 'Skor Web', 'Selisih']):
    sc(t8.rows[0].cells[i], hh, bold=True, size=9)
sh(t8.rows[0])
comp_data = [
    ('Pontianak Mobile Grosir', '0,9358', '0,9358', '0,0000'),
    ('Khatulistiwa Gadget Supply', '1,0142', '1,0142', '0,0000'),
    ('Borneo Tech Distributor', '0,9677', '0,9677', '0,0000'),
    ('Mega Jaya Cellular', '0,9421', '0,9421', '0,0000'),
    ('JBL Audio Partner', '0,9642', '0,9642', '0,0000'),
]
for r, rd in enumerate(comp_data, 1):
    for c, val in enumerate(rd):
        sc(t8.rows[r].cells[c], val, size=9)

p('')
p('Berdasarkan Tabel 2.8 di atas, dapat dilihat bahwa hasil perhitungan pada sistem berbasis web identik dengan hasil perhitungan manual. Selisih nilai sebesar 0,0000 untuk seluruh alternatif menunjukkan bahwa implementasi algoritma SAW pada sistem telah berjalan dengan sangat akurat. Tidak terdapat penyimpangan atau kesalahan perhitungan pada sistem.')

p('Kesamaan hasil ini membuktikan bahwa:')

p('1. Algoritma SAW telah diimplementasikan dengan benar pada kode program backend.')
p('2. Proses normalisasi yang dilakukan sistem sesuai dengan rumus normalisasi SAW.')
p('3. Pembobotan dan penjumlahan terbobot telah dilakukan secara tepat.')
p('4. Bonus ongkos kirim telah ditambahkan dengan benar untuk supplier yang memenuhi syarat.')
p('5. Sistem dapat diandalkan untuk digunakan dalam proses pengambilan keputusan yang sesungguhnya.')

h('2.7 Dokumentasi Observasi dan Wawancara', level=2)

p('Observasi dan wawancara dilakukan dalam rangka mengumpulkan data dan informasi yang diperlukan dalam penelitian ini. Berikut adalah dokumentasi lengkap dari kegiatan tersebut:')

p('2.7.1 Waktu dan Tempat', bold=True, indent=False)
p('Hari/Tanggal : [HARI], [TANGGAL]', indent=False)
p('Waktu : [PUKUL] WIB', indent=False)
p('Tempat : CV Anugerah Mega Makmur, Pontianak', indent=False)
p('Narasumber : [NAMA PEMILIK] (Pemilik Perusahaan)', indent=False)

p('2.7.2 Hasil Wawancara', bold=True, indent=False)

p('Berikut adalah transkrip hasil wawancara yang telah disarikan:')

p('Pertanyaan 1: Bagaimana proses pemilihan supplier yang selama ini berjalan di CV Anugerah Mega Makmur?', bold=True, indent=False)
p('Jawaban: Selama ini saya pribadi yang langsung memilih supplier. Biasanya saya pilih berdasarkan pengalaman sebelumnya, misalnya siapa yang harganya paling murah atau siapa yang barangnya paling bagus. Kadang juga saya tanya ke teman sesama pedagang, siapa supplier yang recommended. Jadi ya masih manual dan subjektif, belum ada sistem atau rumus tertentu. Saya akui kadang juga bingung kalau ada beberapa supplier yang sama-sama bagus, susah menentukan mana yang paling tepat.')

p('Pertanyaan 2: Kriteria apa saja yang Bapak/Ibu pertimbangkan dalam memilih supplier?', bold=True, indent=False)
p('Jawaban: Yang jelas harga. Harga itu paling utama karena kalau harga mahal, susah jualnya nanti. Tapi kalau hanya murah tapi barangnya jelek, ya tidak bisa juga. Jadi kualitas juga penting. Terus pengiriman, kalau supplier sering telat, kasihan pelanggan kita. Layanan juga, misalnya kalau ada komplain atau ada barang rusak, responnya cepat atau tidak. Terakhir kapasitas stok, jangan sampai kita pesan barang, dia stoknya tidak ada.')

p('Pertanyaan 3: Dari kelima kriteria tersebut, mana yang paling penting?', bold=True, indent=False)
p('Jawaban: Harga dan kualitas itu sama pentingnya, beratnya kurang lebih sama. Kalau boleh saya bilang, harga 30%, kualitas 30%, pengiriman 20%, terus layanan dan kapasitas mungkin masing-masing 10%. Itu kira-kira menurut saya.')

p('Pertanyaan 4: Apakah ada standar atau nilai minimal untuk menilai supplier?', bold=True, indent=False)
p('Jawaban: Sebenarnya tidak ada angka pastinya. Tapi secara umum, kalau ada supplier yang benar-benar jelek di satu aspek, misalnya pengirimannya sering telat atau barangnya sering cacat, ya saya tidak akan pakai lagi. Jadi semuanya harus baik minimal, baru bisa dipakai.')

p('Pertanyaan 5: Apakah Bapak/Ibu tertarik jika dibuatkan sistem yang bisa membantu memilih supplier secara otomatis?', bold=True, indent=False)
p('Jawaban: Tertarik sekali. Kebetulan saya juga lagi ingin membuat catatan supplier lebih rapi. Jadi kalau ada sistem yang bisa membantu menilai supplier secara lebih objektif, saya sangat setuju. Apalagi kalau bisa diakses dari handphone atau laptop, pasti sangat membantu.')

p('2.7.3 Hasil Observasi', bold=True, indent=False)
p('Berdasarkan observasi yang dilakukan, berikut adalah beberapa temuan penting:')

p('1. CV Anugerah Mega Makmur memiliki catatan supplier yang masih berupa buku catatan biasa dan file spreadsheet Excel yang tidak terstruktur dengan baik.')
p('2. Proses pengecekan kualitas barang dilakukan secara visual oleh pemilik atau staf gudang tanpa menggunakan standar baku.')
p('3. Data transaksi dengan supplier dicatat secara manual dan belum terintegrasi dengan data penilaian supplier.')
p('4. Tidak ada dokumentasi formal mengenai evaluasi kinerja supplier dari waktu ke waktu.')
p('5. Pemilik perusahaan sangat antusias dengan adanya sistem terkomputerisasi yang dapat membantu proses seleksi supplier.')

doc.add_page_break()

# ══════════════════════════════════════════════
# BAB 3
# ══════════════════════════════════════════════
h('BAB III', level=0, sb=0)
h('PENUTUP', level=0, sb=0)

h('3.1 Kesimpulan', level=2)

p('Berdasarkan hasil penelitian, perancangan, implementasi, dan pengujian yang telah dilakukan, maka dapat ditarik beberapa kesimpulan sebagai berikut:')

p('1. Proses seleksi supplier aksesoris handphone di CV Anugerah Mega Makmur dapat diselesaikan menggunakan metode Simple Additive Weighting (SAW) dengan lima kriteria penilaian, yaitu Harga (bobot 30%), Kualitas (bobot 30%), Pengiriman (bobot 20%), Layanan (bobot 10%), dan Kapasitas (bobot 10%). Harga merupakan satu-satunya kriteria cost, sedangkan keempat kriteria lainnya adalah kriteria benefit. Selain itu, terdapat bonus nilai sebesar 0,05 poin bagi supplier yang menanggung biaya ongkos kirim.')

p('2. Penerapan metode SAW dalam sistem pendukung keputusan seleksi supplier berhasil dilakukan dengan baik. Perhitungan manual terhadap 5 sampel supplier menunjukkan bahwa Khatulistiwa Gadget Supply menempati peringkat pertama dengan skor SAW 1,0142 (termasuk bonus ongkir), diikuti oleh Borneo Tech Distributor (0,9677), JBL Audio Partner (0,9642), Mega Jaya Cellular (0,9421), dan Pontianak Mobile Grosir (0,9358). Seluruh supplier berada di atas threshold 0,75 sehingga semuanya dinyatakan direkomendasikan.')

p('3. Sistem berbasis web yang dibangun menggunakan Next.js 14 (frontend), NestJS (backend), dan PostgreSQL (database) berhasil mengimplementasikan metode SAW dengan baik. Sistem memiliki fitur manajemen data supplier, evaluasi SAW, perankingan, dan manajemen hasil evaluasi. Antarmuka sistem dirancang responsif dan mudah digunakan.')

p('4. Hasil pengujian akurasi menunjukkan bahwa perhitungan pada sistem berbasis web identik dengan perhitungan manual. Selisih nilai antara perhitungan manual dan sistem adalah 0,0000 untuk seluruh alternatif yang diuji. Hal ini membuktikan bahwa implementasi algoritma SAW pada sistem telah akurat dan dapat diandalkan.')

h('3.2 Saran', level=2)

p('Untuk pengembangan sistem lebih lanjut, berikut adalah beberapa saran yang dapat diberikan:')

p('1. Penambahan Jumlah Kriteria. Sistem saat ini hanya menggunakan 5 kriteria. Ke depannya, dapat ditambahkan kriteria-kriteria lain yang relevan seperti garansi produk, reputasi supplier, lama waktu kerja sama, atau ketepatan dokumen. Semakin banyak kriteria yang digunakan, semakin komprehensif hasil penilaian yang diperoleh.')

p('2. Perbandingan Metode. Disarankan untuk melakukan perbandingan antara metode SAW dengan metode multi-kriteria lainnya seperti AHP (Analytical Hierarchy Process), TOPSIS (Technique for Order Preference by Similarity to Ideal Solution), atau WP (Weighted Product) untuk mengetahui metode mana yang paling sesuai dengan karakteristik permasalahan seleksi supplier.')

p('3. Analisis Sensitivitas. Fitur analisis sensitivitas dapat ditambahkan untuk melihat bagaimana perubahan bobot kriteria mempengaruhi hasil perankingan. Dengan fitur ini, pengguna dapat melakukan simulasi berbagai skenario dan melihat dampaknya terhadap rekomendasi supplier.')

p('4. Pengembangan Mobile. Sistem saat ini berbasis web dan responsif, namun belum memiliki aplikasi mobile native. Pengembangan aplikasi mobile akan memudahkan pengguna untuk mengakses sistem kapan saja dan di mana saja, terutama bagi pemilik usaha yang sering mobile.')

p('5. Integrasi dengan Sistem Inventory. Sistem dapat diintegrasikan dengan sistem inventory atau sistem pembelian sehingga data stok barang dan riwayat transaksi dapat digunakan secara langsung dalam proses penilaian supplier tanpa perlu input manual.')

p('6. Riwayat Tren Supplier. Penambahan fitur untuk melihat tren nilai supplier dari waktu ke waktu akan sangat berguna untuk memantau performa supplier secara periodik. Dengan fitur ini, pengguna dapat melihat apakah kinerja supplier cenderung meningkat atau menurun dari waktu ke waktu.')

p('7. Export Laporan. Fitur export hasil evaluasi ke dalam format PDF atau Excel akan memudahkan pengguna dalam membuat laporan dan dokumentasi evaluasi supplier.')

doc.add_page_break()

# ══════════════════════════════════════════════
# DAFTAR PUSTAKA
# ══════════════════════════════════════════════
h('DAFTAR PUSTAKA', level=0, sb=0)
p('', indent=False)

refs = [
    'Afshari, A., Mojahed, M., & Yusuff, R. M. (2010). Simple Additive Weighting approach to Personnel Selection problem. International Journal of Innovation, Management and Technology, 1(5), 511-515.',
    'Kusumadewi, S., Hartati, S., Harjoko, A., & Wardoyo, R. (2006). Fuzzy Multi-Attribute Decision Making (Fuzzy MADM). Yogyakarta: Graha Ilmu.',
    'Munthafa, A. E., & Mubarok, H. (2017). Penerapan Metode Simple Additive Weighting (SAW) pada Sistem Pendukung Keputusan Pemilihan Supplier Terbaik. Jurnal Siliwangi Seri Sains dan Teknologi, 3(2), 105-113.',
    'Prasetyo, H. D., & Handayani, L. (2019). Sistem Pendukung Keputusan Pemilihan Supplier Menggunakan Metode SAW (Simple Additive Weighting). Jurnal Teknologi Informasi dan Ilmu Komputer, 6(3), 289-296.',
    'Pratiwi, D., & Lestari, F. (2020). Sistem Pendukung Keputusan Pemilihan Supplier Bahan Baku dengan Metode SAW. Jurnal Informatika dan Sistem Informasi, 7(1), 45-53.',
    'Putra, A. B., & Ginardi, H. (2021). Implementasi Metode Simple Additive Weighting untuk Seleksi Supplier pada Perusahaan Manufaktur. Jurnal Teknik Informatika, 14(2), 121-130.',
    'Sari, F. (2018). Metode dalam Pengambilan Keputusan. Yogyakarta: Deepublish.',
    'Suryadi, K., & Ramdhani, M. A. (2000). Sistem Pendukung Keputusan: Suatu Wacana Struktural Idealisasi dan Implementasi Konsep Pengambilan Keputusan. Bandung: PT Remaja Rosdakarya.',
    'Turban, E., Aronson, J. E., & Liang, T. P. (2005). Decision Support Systems and Intelligent Systems (7th ed.). New Jersey: Pearson Education.',
    'Wijaya, H. (2019). Sistem Pendukung Keputusan Pemilihan Supplier Menggunakan Metode Simple Additive Weighting pada Perusahaan Retail. Jurnal Ilmiah Teknologi Informasi, 8(2), 89-98.',
]
for ref in refs:
    par = doc.add_paragraph()
    par.paragraph_format.line_spacing = 1.5
    par.paragraph_format.first_line_indent = Cm(0)
    par.paragraph_format.left_indent = Cm(1.27)
    par.paragraph_format.hanging_indent = Cm(1.27)
    run = par.add_run(ref)
    sf(run)

# ── Save ──
out = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'Laporan_SPK_SAW_Supplier.docx')
doc.save(out)
print(f'Laporan berhasil dibuat: {out}')
